from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent
FONT = "AppleGothic"
INK = "20242A"
BLUE = "24577A"
MUTED = "68727D"
LIGHT = "F2F4F7"
ACCENT = "E8EEF5"
YELLOW = "FFF2CC"
WHITE = "FFFFFF"
WIDTH_DXA = 9360
INDENT_DXA = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=INDENT_DXA):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, 8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_document(doc, running_label):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.76)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.36)
    section.footer_distance = Inches(0.36)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for style_name, size, color, before, after in (
        ("Heading 1", 12.5, BLUE, 10, 4),
        ("Heading 2", 11.5, BLUE, 8, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    r = header.add_run(running_label)
    set_font(r, 8.5, bold=True, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.paragraph_format.space_before = Pt(0)
    add_page_field(footer)


def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(title)
    set_font(r, 22, bold=True, color=INK)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(14)
        r2 = p2.add_run(subtitle)
        set_font(r2, 10, color=MUTED)


def add_field_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value, fill in rows:
        cells = table.add_row().cells
        cells[0].text = ""
        cells[1].text = ""
        r1 = cells[0].paragraphs[0].add_run(label)
        set_font(r1, 9.5, bold=True, color=BLUE)
        r2 = cells[1].paragraphs[0].add_run(value)
        set_font(r2, 9.5, color=INK)
        set_cell_shading(cells[0], LIGHT)
        if fill:
            set_cell_shading(cells[1], fill)
    set_table_geometry(table, [2160, 7200])
    return table


def add_clause(doc, number, heading, paragraphs, items=None):
    h = doc.add_paragraph(style="Heading 1")
    h.add_run(f"第{number}条（{heading}）")
    for ptext in paragraphs:
        p = doc.add_paragraph(ptext)
        p.paragraph_format.first_line_indent = Inches(0.22)
    if items:
        for item in items:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.46)
            p.paragraph_format.first_line_indent = Inches(-0.22)
            p.paragraph_format.space_after = Pt(3)
            p.add_run(item)


def build_contract():
    doc = Document()
    configure_document(doc, "東京リーグWebサイトリニューアル｜業務委託契約書")
    add_title(doc, "Webサイト制作業務委託契約書", "東京リーグ Webサイトリニューアル")

    add_field_table(doc, [
        ("委託者（甲）", "【団体正式名称】", YELLOW),
        ("代表者", "【役職・氏名】", YELLOW),
        ("所在地", "【団体所在地】", YELLOW),
        ("受託者（乙）", "【氏名】", YELLOW),
        ("住所", "【住所】", YELLOW),
        ("契約日", "2026年6月19日", None),
    ])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(8)
    p.add_run("甲および乙は、東京リーグWebサイトのリニューアル制作業務について、以下のとおり業務委託契約（以下「本契約」という。）を締結する。")

    add_clause(doc, 1, "目的", [
        "本契約は、甲が乙に委託するWebサイト制作業務の内容、報酬、納品、権利および責任分担を定めることを目的とする。",
        "本契約の規定は、契約締結日前に甲の依頼に基づき乙が実施した本業務にも適用する。",
    ])

    add_clause(doc, 2, "委託業務", [
        "甲は乙に対し、次の業務（以下「本業務」という。）を委託し、乙はこれを受託する。",
    ], [
        "公開Webサイトの情報設計、UI/UX設計、レスポンシブデザインおよび実装",
        "ニュース、大会、リーグ、チーム、試合結果、順位表および資料配布機能の実装",
        "管理者向け管理画面、Google認証、権限管理、担当リーグ割当および監査機能の実装",
        "データベース設計、画像・PDF等のファイル管理、既存データおよび素材の移行支援",
        "セキュリティ対策、動作試験、公開準備および納品確認",
        "仕様書、管理者ツール説明書その他本業務に付随する文書の作成",
    ])

    add_clause(doc, 3, "納品物", [
        "乙は、ソースコード一式、公開サイトおよび管理者ツール、データベース定義、設定手順、仕様書PDFならびに管理者ツール説明書PDFを納品物とする。",
        "秘密情報、個人情報、各種サービスの秘密鍵、アクセストークンおよびパスワードは、納品物に直接記録せず、甲乙間で別途安全な方法により引き渡す。",
    ])

    add_clause(doc, 4, "納期および納品方法", [
        "納期は2026年6月20日とし、乙は甲が指定するリポジトリ、クラウド環境または電磁的方法により納品する。ただし、甲乙協議のうえ書面または電子メールにより変更できる。",
        "本番環境固有のドメイン、外部サービス、認証設定その他甲または第三者の対応を要する事項は、必要情報および権限が提供された後に実施する。",
    ])

    add_clause(doc, 5, "委託料", [
        "本業務の委託料は、金4,800,000円（消費税等を含む契約総額）とする。",
        "ドメイン取得・更新料、外部サービスの有料利用料、サーバー利用料、ストレージ超過料、決済手数料その他第三者へ支払う実費その他本サイトの運用に必要な経費は、本委託料に含まないものとし、甲が直接または乙の請求に基づき実費を負担する。ただし、本契約時点ではサーバーおよびストレージの利用料は無償枠の範囲内を前提とする。",
        "法令上源泉徴収が必要な場合、甲は対象費目について所定額を控除して支払い、控除額および納付内容を乙に通知する。",
    ])

    add_clause(doc, 6, "支払方法", [
        "甲は、乙が発行する請求書に基づき、2026年7月31日までに乙指定の銀行口座へ振り込む。振込手数料は甲の負担とする。",
        "支払条件を変更する場合は、甲乙双方が書面または電子メールで合意する。",
    ])

    add_clause(doc, 7, "検収", [
        "甲は納品後10営業日以内に検査を行い、契約内容に適合しない点がある場合は、その内容を具体的に記載して乙へ通知する。期間内に通知がない場合、納品物は検収に合格したものとみなす。",
        "軽微な表示差異、外部サービスの仕様変更、甲が提供した素材またはデータに起因する事象は、検収不合格の理由としない。",
    ])

    add_clause(doc, 8, "変更および追加作業", [
        "第2条の範囲を超える機能追加、仕様変更、大幅なデザイン変更、データ追加、納品後の運用代行その他の追加作業は、本契約に含まない。",
        "追加作業を行う場合、乙は内容、金額および納期を提示し、甲の承認を得たうえで実施する。",
    ])

    add_clause(doc, 9, "契約不適合への対応", [
        "乙は検収完了日から30日間、本業務の実装に起因する再現可能な不具合を無償で修正する。",
        "甲または第三者による変更、誤操作、外部サービスの障害・仕様変更、対応外環境、想定を超えるアクセスまたはサイバー攻撃に起因する事象は無償修正の対象外とする。",
    ])

    add_clause(doc, 10, "保守運用", [
        "納品後のコンテンツ更新、問い合わせ対応、監視、バックアップ、依存関係更新、障害対応、機能改修その他の保守運用は本契約に含まず、必要に応じて別途契約する。",
    ])

    add_clause(doc, 11, "知的財産権", [
        "乙が本業務において新たに作成した成果物の著作権（著作権法第27条および第28条に定める権利を含む。）は、委託料の全額支払完了時に甲へ移転する。乙は、甲および甲が指定する者に対し著作者人格権を行使しない。",
        "乙または第三者が従前から保有する汎用的な技術、ノウハウ、テンプレート、ライブラリ、オープンソースソフトウェア、フォント、画像その他第三者素材の権利は移転せず、それぞれの利用条件に従う。",
        "乙は、甲の事前承諾なく、非公開情報または個人情報を実績紹介に使用しない。",
    ])

    add_clause(doc, 12, "資料提供および協力", [
        "甲は、本業務に必要な文章、画像、ロゴ、過去データ、各種アカウント情報および確認結果を適時提供し、提供物を利用する正当な権限を有することを保証する。",
        "甲の提供遅延または回答遅延により納期へ影響が生じる場合、乙は合理的な範囲で納期を変更できる。",
    ])

    add_clause(doc, 13, "秘密保持および個人情報", [
        "甲および乙は、本業務に関連して知り得た相手方の非公開情報を、相手方の事前承諾なく第三者に開示または本業務以外に利用しない。ただし、法令に基づく開示および既に公知の情報を除く。",
        "甲および乙は、個人情報および認証情報を必要最小限の範囲で取り扱い、漏えい防止のため合理的な安全管理措置を講じる。",
    ])

    add_clause(doc, 14, "再委託", [
        "乙は、本業務の全部を第三者へ再委託してはならない。本業務の一部を再委託する場合、乙は甲に対して本契約上の責任を負う。",
    ])

    add_clause(doc, 15, "解除", [
        "一方当事者が本契約に重大な違反をし、相当期間を定めた是正要求後も改善しない場合、相手方は本契約を解除できる。",
        "甲の都合により本業務を中途終了する場合、甲は終了時までに完了した業務、確保済み工数および発生実費に相当する金額を乙へ支払う。",
    ])

    add_clause(doc, 16, "損害賠償および免責", [
        "当事者が本契約に関連して相手方へ負う損害賠償責任は、故意または重大な過失がある場合を除き、通常かつ直接の損害に限り、その総額は本契約の委託料を上限とする。",
        "天災、通信障害、外部サービス障害、法令・プラットフォーム仕様の変更その他当事者の合理的な支配を超える事由による履行遅延または不能について、当事者は責任を負わない。",
    ])

    add_clause(doc, 17, "反社会的勢力の排除", [
        "甲および乙は、自らが反社会的勢力に該当せず、反社会的勢力と関係を有しないことを表明し保証する。違反が判明した場合、相手方は催告なく本契約を解除できる。",
    ])

    add_clause(doc, 18, "協議および管轄", [
        "本契約に定めのない事項または解釈上の疑義については、甲乙誠意をもって協議する。",
        "本契約は日本法に準拠し、本契約に関する紛争については東京地方裁判所または東京簡易裁判所を第一審の専属的合意管轄裁判所とする。",
    ])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.add_run("本契約成立の証として、本書2通を作成し甲乙各1通を保有するか、または電磁的記録を作成し、双方が電子署名その他合意した方法により保管する。")

    doc.add_paragraph("2026年6月19日")
    sig = doc.add_table(rows=2, cols=2)
    sig.style = "Table Grid"
    entries = [
        ("甲（委託者）", "団体名：【団体正式名称】\n所在地：【団体所在地】\n代表者：【役職・氏名】　署名／押印：________________"),
        ("乙（受託者）", "氏名：【氏名】\n住所：【住所】\n署名／押印：________________"),
    ]
    for idx, (label, value) in enumerate(entries):
        sig.rows[idx].cells[0].text = ""
        sig.rows[idx].cells[1].text = ""
        r1 = sig.rows[idx].cells[0].paragraphs[0].add_run(label)
        set_font(r1, 9.5, bold=True, color=BLUE)
        r2 = sig.rows[idx].cells[1].paragraphs[0].add_run(value)
        set_font(r2, 9.5)
        set_cell_shading(sig.rows[idx].cells[0], LIGHT)
        set_cell_shading(sig.rows[idx].cells[1], YELLOW)
    set_table_geometry(sig, [1800, 7560])

    path = OUT / "tokyo-league-web-renewal-contract.docx"
    doc.save(path)
    return path


def add_invoice_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["項目", "数量", "単価", "金額"]
    for idx, value in enumerate(headers):
        p = table.rows[0].cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(value)
        set_font(r, 9, bold=True, color=WHITE)
        set_cell_shading(table.rows[0].cells[idx], BLUE)
    set_repeat_table_header(table.rows[0])
    for label, qty, unit, amount in rows:
        cells = table.add_row().cells
        values = [label, qty, unit, amount]
        for idx, value in enumerate(values):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(value)
            set_font(r, 9)
    set_table_geometry(table, [5220, 900, 1530, 1710])
    return table


def build_invoice():
    doc = Document()
    configure_document(doc, "東京リーグWebサイトリニューアル｜請求書")
    add_title(doc, "請 求 書")

    meta = doc.add_table(rows=1, cols=2)
    meta.style = "Table Grid"
    left = meta.rows[0].cells[0]
    right = meta.rows[0].cells[1]
    left.text = ""
    right.text = ""
    p = left.paragraphs[0]
    r = p.add_run("【団体正式名称】 御中")
    set_font(r, 13, bold=True)
    p.add_run("\n【担当部署・担当者名】").font.size = Pt(9.5)
    set_cell_shading(left, YELLOW)
    lines = [
        ("請求書番号", "TL-20260619-001"),
        ("発行日", "2026年6月19日"),
        ("支払期限", "2026年7月31日"),
    ]
    for label, value in lines:
        p = right.add_paragraph() if right.paragraphs[0].text else right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr = p.add_run(f"{label}：{value}")
        set_font(rr, 9.5)
    set_table_geometry(meta, [5040, 4320])

    issuer = doc.add_paragraph()
    issuer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    issuer.paragraph_format.space_before = Pt(8)
    issuer.paragraph_format.space_after = Pt(10)
    rr = issuer.add_run("請求者：【氏名】\n住所：【住所】\n電話・メール：【連絡先】\n適格請求書発行事業者登録番号：【登録済みの場合のみ記載】")
    set_font(rr, 9.5)

    lead = doc.add_paragraph("下記のとおりご請求申し上げます。")
    lead.paragraph_format.space_after = Pt(6)

    total_box = doc.add_table(rows=1, cols=2)
    total_box.style = "Table Grid"
    total_box.rows[0].cells[0].text = ""
    total_box.rows[0].cells[1].text = ""
    r1 = total_box.rows[0].cells[0].paragraphs[0].add_run("ご請求金額")
    set_font(r1, 12, bold=True, color=WHITE)
    set_cell_shading(total_box.rows[0].cells[0], BLUE)
    p2 = total_box.rows[0].cells[1].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run("4,800,000 円")
    set_font(r2, 17, bold=True)
    set_cell_shading(total_box.rows[0].cells[1], ACCENT)
    set_table_geometry(total_box, [2520, 6840])

    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    rows = [
        ("要件整理・サイト設計", "1式", "350,000円", "350,000円"),
        ("UI/UX・レスポンシブデザイン", "1式", "550,000円", "550,000円"),
        ("公開サイト開発", "1式", "650,000円", "650,000円"),
        ("管理画面・CMS開発", "1式", "1,200,000円", "1,200,000円"),
        ("DB・認証・権限・ファイル管理", "1式", "750,000円", "750,000円"),
        ("既存データ・画像移行", "1式", "300,000円", "300,000円"),
        ("セキュリティ・テスト・公開対応", "1式", "450,000円", "450,000円"),
        ("仕様書・管理者マニュアル", "1式", "300,000円", "300,000円"),
        ("進行管理・修正対応", "1式", "250,000円", "250,000円"),
    ]
    add_invoice_table(doc, rows)

    totals = doc.add_table(rows=4, cols=2)
    totals.style = "Table Grid"
    summary = [
        ("契約金額（消費税等を含む総額）", "4,800,000円"),
        ("源泉徴収額", "法令上必要な場合は支払者側で計算・控除"),
        ("控除後振込額", "控除がある場合は支払者から通知"),
        ("請求総額", "4,800,000円"),
    ]
    for idx, (label, value) in enumerate(summary):
        totals.rows[idx].cells[0].text = ""
        totals.rows[idx].cells[1].text = ""
        p1 = totals.rows[idx].cells[0].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr1 = p1.add_run(label)
        set_font(rr1, 9, bold=idx == 3)
        p2 = totals.rows[idx].cells[1].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr2 = p2.add_run(value)
        set_font(rr2, 9, bold=idx == 3)
        if idx == 3:
            set_cell_shading(totals.rows[idx].cells[0], ACCENT)
            set_cell_shading(totals.rows[idx].cells[1], ACCENT)
    set_table_geometry(totals, [5940, 3420])

    h = doc.add_paragraph(style="Heading 1")
    h.add_run("お振込先")
    bank = doc.add_table(rows=4, cols=2)
    bank.style = "Table Grid"
    bank_rows = [
        ("金融機関・支店", "【銀行名・支店名】"),
        ("口座種別・番号", "【普通／当座・口座番号】"),
        ("口座名義", "【口座名義（カナ）】"),
        ("振込手数料", "ご負担をお願いいたします"),
    ]
    for idx, (label, value) in enumerate(bank_rows):
        bank.rows[idx].cells[0].text = ""
        bank.rows[idx].cells[1].text = ""
        r1 = bank.rows[idx].cells[0].paragraphs[0].add_run(label)
        set_font(r1, 9, bold=True, color=BLUE)
        r2 = bank.rows[idx].cells[1].paragraphs[0].add_run(value)
        set_font(r2, 9)
        set_cell_shading(bank.rows[idx].cells[0], LIGHT)
        if idx < 3:
            set_cell_shading(bank.rows[idx].cells[1], YELLOW)
    set_table_geometry(bank, [2520, 6840])

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    note.paragraph_format.space_after = Pt(0)
    rr = note.add_run("備考：本請求は、東京リーグWebサイトリニューアル制作業務委託契約に基づくものです。源泉徴収を行う場合は、対象費目、控除額および納付内容をご通知ください。")
    set_font(rr, 8.8, color=MUTED)

    path = OUT / "tokyo-league-web-renewal-invoice.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    print(build_contract())
    print(build_invoice())
